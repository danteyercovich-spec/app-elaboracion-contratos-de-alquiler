"""
AutoContract - Backend FastAPI
Gestor Legal Conversacional para Contratos de Alquiler
Soporta: OpenAI (gpt-4o-mini) y Anthropic Claude
"""

import os
import re
import json
import io
import traceback
from typing import Any
from pathlib import Path
from fastapi import FastAPI, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

# Imports opcionales de proveedores
try:
    import anthropic
except ImportError:
    anthropic = None

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

# Cargar .env con ruta absoluta basada en la ubicación de este archivo (backend/)
BASE_DIR = Path(__file__).resolve().parent
ENV_PATH = BASE_DIR / ".env"
load_dotenv(dotenv_path=ENV_PATH, override=True)


# ─── Configuración unificada de proveedores ───────────────────────────────────
AI_PROVIDER = os.getenv("AI_PROVIDER", "openai").lower().strip()

# Variables de key compartidas para logs
USED_KEY_NAME = "NONE"
CLEAN_API_KEY = None
claude_client = None
openai_client = None
CLAUDE_MODEL = "claude-3-5-sonnet-20240620"
OPENAI_MODEL = "gpt-4o-mini"

def inicializar_clientes():
    global claude_client, openai_client, USED_KEY_NAME, CLEAN_API_KEY, CLAUDE_MODEL, OPENAI_MODEL
    
    if AI_PROVIDER == "claude":
        if not anthropic:
            raise RuntimeError("Instale anthropic: pip install anthropic")
        
        USED_KEY_NAME = "ANTHROPIC_API_KEY"
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if not api_key:
            USED_KEY_NAME = "CLAUDE_API_KEY"
            api_key = os.getenv("CLAUDE_API_KEY")
        
        CLEAN_API_KEY = api_key.strip() if api_key else None
        CLAUDE_MODEL = os.getenv("CLAUDE_MODEL", "claude-3-5-sonnet-20240620").strip()
        
        claude_client = anthropic.Anthropic(api_key=CLEAN_API_KEY)
        openai_client = None
        print(f"[OK] Cliente Claude inicializado ({CLAUDE_MODEL})")
    else:
        if not OpenAI:
            raise RuntimeError("Instale openai: pip install openai")
            
        openai_key = os.getenv("OPENAI_API_KEY")
        CLEAN_API_KEY = openai_key.strip() if openai_key else None
        USED_KEY_NAME = "OPENAI_API_KEY"
        
        openai_client = OpenAI(api_key=CLEAN_API_KEY)
        OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini").strip()
        claude_client = None
        print(f"[OK] Cliente OpenAI inicializado ({OPENAI_MODEL})")

# Inicialización al arranque
inicializar_clientes()


# ─── Función unificada de llamada a IA ───────────────────────────────────────

def llamar_ia(system_prompt: str, user_message: str,
              messages_history: list = None,
              json_mode: bool = False,
              temperature: float = 0.2) -> str:
    """
    Llama al proveedor configurado (OpenAI o Claude) y devuelve el texto.
    """
    if AI_PROVIDER == "claude":
        # ── Claude ──
        msgs = []
        if messages_history:
            for m in messages_history:
                msgs.append({"role": m["role"], "content": m["content"]})
        else:
            msgs.append({"role": "user", "content": user_message})

        if json_mode:
            system_prompt += "\n\nIMPORTANTE: Responde ÚNICAMENTE con JSON válido, sin texto adicional."

        if not claude_client:
             raise HTTPException(status_code=500, detail="El cliente de Claude no ha sido inicializado. Verifique su API Key.")

        try:
            response = claude_client.messages.create(
                model=CLAUDE_MODEL,
                max_tokens=8192,
                system=system_prompt,
                messages=msgs,
                temperature=temperature
            )
            return response.content[0].text
        except Exception as e:
            # Manejo dinámico de excepciones de Anthropic sin depender del import estático
            err_type = type(e).__name__
            if err_type == 'NotFoundError':
                raise HTTPException(
                    status_code=404, 
                    detail=f"Modelo '{CLAUDE_MODEL}' no encontrado o no habilitado para esta API Key. "
                           "Verifique CLAUDE_MODEL en su .env."
                )
            elif err_type == 'AuthenticationError':
                raise HTTPException(
                    status_code=401,
                    detail="Error de autenticación con Anthropic. Verifique su API Key en el archivo .env."
                )
            raise HTTPException(status_code=500, detail=f"Error en llamada a Claude: {str(e)}")

    else:
        # ── OpenAI ──
        msgs = [{"role": "system", "content": system_prompt}]
        if messages_history:
            msgs.extend(messages_history)
        else:
            msgs.append({"role": "user", "content": user_message})

        kwargs = {
            "model": OPENAI_MODEL,
            "messages": msgs,
            "temperature": temperature,
        }
        if json_mode:
            kwargs["response_format"] = {"type": "json_object"}

        response = openai_client.chat.completions.create(**kwargs)
        return response.choices[0].message.content


def parsear_json(texto: str) -> dict:
    """Parsea JSON de la respuesta, tolerando texto alrededor."""
    texto = texto.strip()
    # Intentar directo
    try:
        return json.loads(texto)
    except json.JSONDecodeError:
        pass
    # Buscar bloque JSON entre ```
    match = re.search(r'```(?:json)?\s*([\s\S]*?)```', texto)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass
    # Buscar primer { ... }
    match = re.search(r'(\{[\s\S]*\})', texto)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError as e:
            raise ValueError(f"JSON malformado o incompleto. Asegúrese de que el contrato no sea demasiado extenso. Error: {str(e)}")
    
    raise ValueError(f"No se pudo encontrar un JSON válido en la respuesta de la IA. Respuesta parcial: {texto[:200]}...")


# ─── FastAPI ──────────────────────────────────────────────────────────────────

app = FastAPI(title="AutoContract API", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

frontend_path = os.path.join(os.path.dirname(__file__), "..", "frontend")
if os.path.exists(frontend_path):
    app.mount("/static", StaticFiles(directory=frontend_path), name="static")


# ─── Modelos Pydantic ────────────────────────────────────────────────────────

class AnalyzeRequest(BaseModel):
    contract_text: str

class AnalyzeResponse(BaseModel):
    variables: list[dict]
    analysis_notes: str

class ChatMessage(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    messages: list[ChatMessage]
    variables: list[dict]
    collected_data: dict[str, Any]

class ChatResponse(BaseModel):
    reply: str
    collected_data: dict[str, Any]
    is_complete: bool
    next_variable: str | None = None

class GenerateRequest(BaseModel):
    contract_template: str
    variables: list[dict]
    collected_data: dict[str, Any]

class GenerateResponse(BaseModel):
    contract_preview: str
    variables_applied: int


# ─── Endpoints ───────────────────────────────────────────────────────────────

# Los archivos estáticos se manejan al final del archivo con app.mount

@app.get("/api/info")
async def info():
    """Devuelve el proveedor de IA activo."""
    return {
        "provider": AI_PROVIDER,
        "model": CLAUDE_MODEL if AI_PROVIDER == "claude" else OPENAI_MODEL
    }


@app.post("/api/analyze", response_model=AnalyzeResponse)
def analyze_contract(request: AnalyzeRequest):
    """
    Analiza el texto del contrato y detecta variables a completar.
    """
    # Logs seguros de diagnóstico en cada llamada
    masked_key = f"{CLEAN_API_KEY[:6]}...{CLEAN_API_KEY[-4:]}" if CLEAN_API_KEY and len(CLEAN_API_KEY) > 10 else "N/A"
    print(f"\n[DIAGNOSTIC] Llamada a /api/analyze")
    print(f" - Ruta .env: {ENV_PATH}")
    print(f" - Proveedor: {AI_PROVIDER}")
    print(f" - Variable: {USED_KEY_NAME}")
    print(f" - Key (mask): {masked_key}")
    print(f" - Modelo: {CLAUDE_MODEL if AI_PROVIDER == 'claude' else OPENAI_MODEL}")

    if not request.contract_text or len(request.contract_text.strip()) < 50:
        raise HTTPException(status_code=400, detail="El texto del contrato es demasiado corto.")

    system_prompt = """Eres un experto legal argentino. Identifica variables en contratos de alquiler.

BUSCA: Locador, Locatario, Garante, Fiador, DNI, CUIT, Domicilios, Montos, Fechas.
REGLA: El "placeholder_text" debe ser el fragmento EXACTO del contrato (ej: ".........." o "DNI N° .....").
IMPORTANTE: Revisa el FINAL del contrato para los GARANTES.

Responde ÚNICAMENTE con este formato JSON:
{
  "variables": [
    {"key": "dniGarante", "label": "DNI del Garante", "placeholder_text": "D.N.I. ....", "type": "dni"}
  ],
  "analysis_notes": "Análisis rápido"
}"""

    try:
        raw = llamar_ia(
            system_prompt=system_prompt,
            user_message=f"Analiza este contrato completo:\n\n{request.contract_text}",
            json_mode=True,
            temperature=0.1
        )
        result = parsear_json(raw)
        variables = result.get("variables", [])

        clean_vars = []
        seen_keys = set()
        for v in variables:
            if isinstance(v, dict) and "key" in v and "label" in v:
                key = v["key"]
                if key not in seen_keys:
                    seen_keys.add(key)
                    clean_vars.append({
                        "key": key,
                        "label": v.get("label", key),
                        "placeholder_text": v.get("placeholder_text", f"{{{{{key.upper()}}}}}"),
                        "type": v.get("type", "texto"),
                        "description": v.get("description", ""),
                        "example": v.get("example", "")
                    })

        return AnalyzeResponse(
            variables=clean_vars,
            analysis_notes=result.get("analysis_notes", "Análisis completado.")
        )

    except HTTPException as e:
        # Re-lanzar HTTPExceptions (como el 401/404 que ya manejamos en llamar_ia)
        raise e
    except Exception as e:
        print("\n[!!!] ERROR CRÍTICO EN /api/analyze")
        print(f"Tipo de error: {type(e).__name__}")
        
        # Si el error tiene atributos de respuesta (como los de Anthropic/httpx)
        if hasattr(e, 'status_code'):
            print(f"Código HTTP real: {e.status_code}")
        if hasattr(e, 'message'):
            print(f"Mensaje exacto: {e.message}")
        
        # Imprimir traceback completo para diagnóstico
        traceback.print_exc()
        
        raise HTTPException(status_code=500, detail=f"Error en el análisis: {str(e)}")


@app.post("/api/chat", response_model=ChatResponse)
def chat(request: ChatRequest):
    """
    Conversación guiada para recopilar datos de las variables.
    """
    variables = request.variables
    collected = request.collected_data.copy()
    pending_vars = [v for v in variables if v["key"] not in collected or not collected[v["key"]]]

    system_prompt = f"""Eres AsistenteContrato, un asistente legal formal para completar contratos de alquiler en Argentina.
Tu único objetivo es preguntarle al usuario CADA UNA de las variables pendientes.

LISTA DE VARIABLES (TODAS DEBEN SER COMPLETADAS):
{json.dumps(variables, ensure_ascii=False, indent=2)}

DATOS ACTUALES:
{json.dumps(collected, ensure_ascii=False, indent=2)}

INSTRUCCIONES:
1. NO des por terminada la entrevista hasta que TODAS las variables tengan un valor.
2. Haz UNA pregunta clara a la vez.
3. Si el usuario da un dato que parece ser para otra variable, extráelo igual.
4. Si detectas un error de formato (ej: un DNI de 3 números), pide corregirlo amablemente.
5. NO te saltes a los Garantes/Fiadores si están en la lista.

Responde con JSON válido:
{{
  "reply": "Tu mensaje al usuario",
  "extracted_data": {{"key_de_la_variable": "valor_extraido"}},
  "is_complete": false,
  "next_variable_key": "key_de_la_siguiente"
}}"""

    try:
        history = [{"role": m.role, "content": m.content} for m in request.messages]

        raw = llamar_ia(
            system_prompt=system_prompt,
            user_message="",
            messages_history=history,
            json_mode=True,
            temperature=0.3
        )
        result = parsear_json(raw)

        extracted = result.get("extracted_data", {})
        collected.update(extracted)

        is_complete = result.get("is_complete", False)
        if not is_complete:
            still_pending = [v for v in variables if v["key"] not in collected or not collected[v["key"]]]
            is_complete = len(still_pending) == 0

        return ChatResponse(
            reply=result.get("reply", "¿Podría repetir ese dato?"),
            collected_data=collected,
            is_complete=is_complete,
            next_variable=result.get("next_variable_key")
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error en el chat: {str(e)}")


@app.post("/api/generate", response_model=GenerateResponse)
def generate_contract(request: GenerateRequest):
    """
    Inyecta los datos en la plantilla. Sustitución en 3 capas para máxima confiabilidad.
    """
    contract = request.contract_template
    applied = 0
    no_match = []

    print(f"\n[GENERATE] Variables recibidas: {len(request.variables)}")
    print(f"[GENERATE] Datos recolectados : {list(request.collected_data.keys())}")

    for variable in request.variables:
        key         = variable["key"]
        placeholder = variable.get("placeholder_text", "").strip()
        value       = str(request.collected_data.get(key, "")).strip()

        if not value:
            print(f"[GENERATE] SKIP  '{key}' — sin valor ingresado")
            continue
        if not placeholder:
            print(f"[GENERATE] SKIP  '{key}' — sin placeholder_text")
            continue

        original_contract = contract
        replaced = False

        # ── CAPA 1: Coincidencia exacta (todas las ocurrencias) ───────────────
        if placeholder in contract:
            contract = contract.replace(placeholder, value)
            replaced = True

        # ── CAPA 2: Normalización de espacios (tabs → espacio, múltiples → uno)
        if not replaced:
            ph_norm  = re.sub(r'\s+', ' ', placeholder)
            ctx_norm = re.sub(r'\s+', ' ', contract)
            if ph_norm in ctx_norm:
                contract = re.sub(r'\s+', ' ', contract)
                contract = contract.replace(ph_norm, value)
                replaced = True

        # ── CAPA 3: Regex flexible (puntos, guiones, espacios variables) ──────
        if not replaced:
            try:
                p_esc  = re.escape(placeholder)
                # Secuencias de puntos → cualquier cantidad de puntos/guiones/espacios
                p_flex = re.sub(r'(\\\\.)+', r'[.\\-\\s]+', p_esc)
                # Guiones bajos flexibles
                p_flex = re.sub(r'\\_+', r'[_\\s]+', p_flex)
                # Espacios flexibles
                p_flex = re.sub(r'\\ ', r'\\s*', p_flex)

                if re.search(p_flex, contract):
                    contract = re.sub(p_flex, value, contract)   # reemplaza TODAS
                    replaced = True
            except re.error as e:
                print(f"[GENERATE] REGEX ERR '{key}': {e}")

        # ── Resultado por variable ────────────────────────────────────────────
        if replaced:
            applied += 1
            print(f"[GENERATE] OK    '{key}' → '{value[:30]}{'...' if len(value)>30 else ''}'")
        else:
            no_match.append(key)
            print(f"[GENERATE] MISS  '{key}' — placeholder no encontrado: '{placeholder[:60]}'")

    print(f"\n[GENERATE] Resultado: {applied}/{len(request.variables)} variables aplicadas")
    if no_match:
        print(f"[GENERATE] Sin match: {no_match}")

    return GenerateResponse(
        contract_preview=contract,
        variables_applied=applied
    )



@app.post("/api/export-docx")
def export_docx(request: GenerateRequest):
    """
    Genera el contrato como .DOCX profesional.
    """
    gen_response = generate_contract(request)
    contract_text = gen_response.contract_preview

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    for para_text in contract_text.split('\n'):
        para_text = para_text.strip()
        if not para_text:
            doc.add_paragraph()
            continue

        is_title = (
            para_text.isupper() or
            re.match(r'^(CLÁUSULA|ARTÍCULO|Cláusula|Artículo|TÍTULO|Título)\s+\w+', para_text) or
            re.match(r'^\d+[\.\-]\s+[A-Z]', para_text) or
            (len(para_text) < 80 and para_text.endswith(':'))
        )

        para = doc.add_paragraph()
        run = para.add_run(para_text)

        if is_title:
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if para_text.isupper() else WD_ALIGN_PARAGRAPH.LEFT
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Inches(0.3)

        para.paragraph_format.space_after = Pt(6)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=contrato_alquiler.docx"}
    )


@app.get("/")
async def root():
    index_path = os.path.join(frontend_path, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    return {"message": "AutoContract API activa", "provider": AI_PROVIDER}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
