import sys
import os
import io
from docx import Document

# Add backend to path to import functions
sys.path.append(os.path.join(os.getcwd(), 'v2', 'backend'))
import main

def create_test_docx():
    doc = Document()
    p = doc.add_paragraph("Esto es una prueba con ")
    # Manually create split runs for {{TEST_FIELD}}
    p.add_run("{{")
    p.add_run("TEST")
    p.add_run("_FIELD")
    p.add_run("}}")
    p.add_run(" y otro campo ")
    run = p.add_run("{{OTRO}}")
    run.bold = True
    
    p2 = doc.add_paragraph("Campo repetido: {{TEST_FIELD}}")
    
    doc.save("test_template.docx")
    print("Plantilla de prueba creada: test_template.docx")
    return "test_template.docx"

def test_logic():
    path = create_test_docx()
    doc = Document(path)
    
    # 1. Test extraction
    placeholders = main.extract_placeholders(doc)
    print(f"Placeholders detectados: {placeholders}")
    assert "TEST_FIELD" in placeholders
    assert "OTRO" in placeholders
    
    # 2. Test replacement
    replacements = {
        "TEST_FIELD": "VALOR_EXITOSO",
        "OTRO": "NEGRILLA"
    }
    doc = main.replace_placeholders(doc, replacements)
    
    # Verify replacements
    full_text = ""
    for p in doc.paragraphs:
        full_text += p.text + "\n"
    
    print(f"Texto resultante:\n{full_text}")
    assert "VALOR_EXITOSO" in full_text
    assert "NEGRILLA" in full_text
    assert "{{" not in full_text
    
    # Verify formatting of "OTRO"
    found_negrilla = False
    for p in doc.paragraphs:
        for r in p.runs:
            if "NEGRILLA" in r.text and r.bold:
                found_negrilla = True
    assert found_negrilla, "El formato negrita se perdió"
    
    print("✅ Pruebas de lógica básica PASADAS")

if __name__ == "__main__":
    try:
        test_logic()
    finally:
        if os.path.exists("test_template.docx"):
            os.remove("test_template.docx")
