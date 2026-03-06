"""
Genera una plantilla .docx de prueba con placeholders {{...}} para validar AutoContract V2.
Ejecutar: python crear_plantilla.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

def h(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def body(text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

# Título
h("CONTRATO DE LOCACIÓN DE VIVIENDA", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# Encabezado
body(
    "En la ciudad de {{CIUDAD}}, a los {{FECHA_DIA}} días del mes de {{FECHA_MES}} "
    "del año {{FECHA_ANIO}}, entre las partes que a continuación se individualizan:"
)
doc.add_paragraph()

# Sección LOCADOR
h("LOCADOR", bold=True)
body(
    "Nombre y Apellido: {{LOCADOR_NOMBRE}}\n"
    "DNI N°: {{LOCADOR_DNI}}\n"
    "CUIT/CUIL: {{LOCADOR_CUIT}}\n"
    "Domicilio legal: {{LOCADOR_DOMICILIO}}"
)
doc.add_paragraph()

# Sección LOCATARIO
h("LOCATARIO/S", bold=True)
body(
    "Nombre y Apellido: {{LOCATARIO_NOMBRE}}\n"
    "DNI N°: {{LOCATARIO_DNI}}\n"
    "CUIT/CUIL: {{LOCATARIO_CUIT}}\n"
    "Domicilio legal: {{LOCATARIO_DOMICILIO}}"
)
doc.add_paragraph()

# Sección GARANTE (campo opcional)
h("GARANTE (opcional)", bold=True)
body(
    "Nombre y Apellido: {{GARANTE_NOMBRE}}\n"
    "DNI N°: {{GARANTE_DNI}}"
)
doc.add_paragraph()

# Cláusulas en tabla
h("CLÁUSULAS", bold=True)
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'

rows_data = [
    ("INMUEBLE",         "El inmueble objeto del presente contrato se encuentra ubicado en {{INMUEBLE_DIRECCION}}, {{INMUEBLE_CIUDAD}}."),
    ("DESTINO",          "El inmueble se destina exclusivamente a {{DESTINO_USO}} (vivienda familiar del locatario)."),
    ("PLAZO",            "El plazo de la locación es de {{PLAZO_MESES}} meses, comenzando el {{FECHA_INICIO}} y finalizando el {{FECHA_FIN}}."),
    ("CANON",            "El canon locativo mensual es de PESOS {{MONTO_ALQUILER}} (${{MONTO_ALQUILER_NUM}})."),
    ("DEPÓSITO",         "Se establece un depósito en garantía equivalente a {{DEPOSITO_MESES}} mes/es de alquiler, por un total de PESOS {{DEPOSITO_MONTO}}."),
    ("ACTUALIZACIÓN",    "El alquiler se actualizará cada {{ACTUALIZACION_PERIODO}} conforme al índice {{ACTUALIZACION_INDICE}}."),
]

for i, (titulo, contenido) in enumerate(rows_data):
    row = table.rows[i]
    row.cells[0].text = titulo
    p = row.cells[0].paragraphs[0]
    p.runs[0].bold = True
    row.cells[1].text = contenido

doc.add_paragraph()

# Firma
body(
    "En prueba de conformidad, las partes firman el presente contrato en {{CANTIDAD_EJEMPLARES}} "
    "ejemplares de un mismo tenor y a un solo efecto, en el lugar y fecha indicados en el encabezamiento."
)
doc.add_paragraph()
doc.add_paragraph()

# Líneas de firma
firma_table = doc.add_table(rows=2, cols=2)
firma_table.rows[0].cells[0].text = "________________________"
firma_table.rows[0].cells[1].text = "________________________"
firma_table.rows[1].cells[0].text = "LOCADOR: {{LOCADOR_NOMBRE}}"
firma_table.rows[1].cells[1].text = "LOCATARIO: {{LOCATARIO_NOMBRE}}"

# Guardar
out = os.path.join(os.path.dirname(__file__), "plantilla_prueba.docx")
doc.save(out)
print(f"Plantilla creada: {out}")
print("Placeholders incluidos:")
placeholders = [
    "CIUDAD", "FECHA_DIA", "FECHA_MES", "FECHA_ANIO",
    "LOCADOR_NOMBRE", "LOCADOR_DNI", "LOCADOR_CUIT", "LOCADOR_DOMICILIO",
    "LOCATARIO_NOMBRE", "LOCATARIO_DNI", "LOCATARIO_CUIT", "LOCATARIO_DOMICILIO",
    "GARANTE_NOMBRE", "GARANTE_DNI",
    "INMUEBLE_DIRECCION", "INMUEBLE_CIUDAD", "DESTINO_USO",
    "PLAZO_MESES", "FECHA_INICIO", "FECHA_FIN",
    "MONTO_ALQUILER", "MONTO_ALQUILER_NUM",
    "DEPOSITO_MESES", "DEPOSITO_MONTO",
    "ACTUALIZACION_PERIODO", "ACTUALIZACION_INDICE",
    "CANTIDAD_EJEMPLARES",
]
for p in placeholders:
    print(f"  {{{{ {p} }}}}")
