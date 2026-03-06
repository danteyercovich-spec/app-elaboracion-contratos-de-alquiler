"""
AutoContract V2 — Backend FastAPI
Flujo determinístico: plantilla .docx con {{PLACEHOLDERS}} → formulario → .docx final
Sin IA en generación.
"""

import re
import io
import os
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from docx.oxml.ns import qn
from typing import Optional
import copy
import traceback

app = FastAPI(title="AutoContract V2")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Exception Handlers ──────────────────────────────────────────────────────

@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    print(f"[ERROR] HTTP {exc.status_code}: {exc.detail}")
    return JSONResponse(
        status_code=exc.status_code,
        content={"error": "HTTP Error", "detail": str(exc.detail)},
    )

@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    error_msg = f"Error inesperado: {str(exc)}"
    print(f"[CRITICAL] {error_msg}")
    traceback.print_exc()
    return JSONResponse(
        status_code=500,
        content={"error": "Internal Server Error", "detail": error_msg},
    )

# ─── Modelos ──────────────────────────────────────────────────────────────────

class GenerateRequest(BaseModel):
    values: dict[str, str]          # { "LOCADOR_NOMBRE": "Juan García", ... }
    optional_empty: list[str] = []  # placeholders marcados como vacíos opcionalmente


# ─── Utilidades docx ─────────────────────────────────────────────────────────

PLACEHOLDER_RE = re.compile(r'\{\{([A-Z0-9_]+)\}\}')


def _replace_in_paragraph(para, replacements: dict):
    """
    Reemplaza placeholders en un párrafo preservando el formato original.
    Verificable mediante log en consola.
    """
    if not replacements:
        return

    full_text = "".join(r.text for r in para.runs)
    matches = list(PLACEHOLDER_RE.finditer(full_text))
    if not matches:
        return

    # Log de verificación pedido por el usuario
    print("V2 DOCX REPLACE RUN-LEVEL v1")

    # Procesar de atrás hacia adelante para no invalidar offsets de texto
    for match in reversed(matches):
        placeholder_name = match.group(1)
        if placeholder_name not in replacements:
            continue
        
        replacement_text = replacements[placeholder_name]
        start_idx, end_idx = match.span()

        # Identificar runs afectados
        affected_indices = []
        current_pos = 0
        for i, run in enumerate(para.runs):
            run_len = len(run.text)
            if current_pos < end_idx and current_pos + run_len > start_idx:
                affected_indices.append(i)
            current_pos += run_len

        if not affected_indices:
            continue

        first_idx = affected_indices[0]
        last_idx = affected_indices[-1]
        
        first_run = para.runs[first_idx]
        last_run = para.runs[last_idx]

        # Calcular offsets relativos al inicio de sus respectivos runs
        pos_before_first = sum(len(para.runs[i].text) for i in range(first_idx))
        rel_start = start_idx - pos_before_first
        
        pos_before_last = sum(len(para.runs[i].text) for i in range(last_idx))
        rel_end = end_idx - pos_before_last

        prefix = first_run.text[:rel_start]
        suffix = last_run.text[rel_end:]

        if first_idx == last_idx:
            # Caso simple: un solo run (el estilo se mantiene naturalmente)
            first_run.text = prefix + replacement_text + suffix
        else:
            # Caso complejo: el estilo se hereda del first_run
            first_run.text = prefix + replacement_text
            last_run.text = suffix
            
            # Eliminar runs intermedios físicamente del XML
            # Recolectamos elementos primero para evitar problemas con la lista dinámica runs
            elements_to_remove = [para.runs[i]._element for i in range(first_idx + 1, last_idx)]
            for el in elements_to_remove:
                if el.getparent() is not None:
                    el.getparent().remove(el)


def extract_placeholders(doc: Document) -> list[str]:
    """Extrae todos los {{PLACEHOLDERS}} únicos del documento en orden de aparición."""
    seen = set()
    ordered = []

    def _scan_text(text: str):
        for match in PLACEHOLDER_RE.finditer(text):
            name = match.group(1)
            if name not in seen:
                seen.add(name)
                ordered.append(name)

    # Párrafos del cuerpo principal
    for para in doc.paragraphs:
        _scan_text(''.join(r.text for r in para.runs))

    # Tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan_text(''.join(r.text for r in para.runs))

    # Headers y footers
    for section in doc.sections:
        for hf in [section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer]:
            if hf is not None:
                for para in hf.paragraphs:
                    _scan_text(''.join(r.text for r in para.runs))

    return ordered


def replace_placeholders(doc: Document, replacements: dict) -> Document:
    """Reemplaza placeholders en el documento conservando formato."""

    def process_para(para):
        _replace_in_paragraph(para, replacements)

    # Body
    for para in doc.paragraphs:
        process_para(para)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_para(para)

    # Headers / Footers
    for section in doc.sections:
        for hf in [section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer]:
            if hf is not None:
                for para in hf.paragraphs:
                    process_para(para)

    return doc


def check_for_remaining_placeholders(doc: Document) -> list[str]:
    """Busca si quedaron placeholders sin reemplazar en el documento final."""
    remaining = []

    def _scan(text):
        for match in PLACEHOLDER_RE.finditer(text):
            remaining.append(match.group(0))

    for para in doc.paragraphs:
        _scan("".join(r.text for r in para.runs))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan("".join(r.text for r in para.runs))

    return list(set(remaining))


# ─── Endpoints ────────────────────────────────────────────────────────────────

# Almacenamiento temporal en memoria de la plantilla cargada (MVP)
_template_store: dict = {}


@app.post("/api/extract")
async def extract(file: UploadFile = File(...)):
    """
    Recibe un .docx, extrae {{PLACEHOLDERS}} únicos y los devuelve.
    También guarda la plantilla en memoria para la generación posterior.
    """
    if not file.filename.endswith('.docx'):
        raise HTTPException(400, detail="Solo se aceptan archivos .docx")

    content = await file.read()

    try:
        doc = Document(io.BytesIO(content))
    except Exception as e:
        print(f"[EXTRACT] Error de lectura docx: {e}")
        raise HTTPException(400, detail=f"No se pudo leer el archivo .docx: {e}")

    placeholders = extract_placeholders(doc)

    if not placeholders:
        print(f"[EXTRACT] Documento sin placeholders: {file.filename}")
        raise HTTPException(422, detail="El documento no contiene placeholders {{...}}. "
                                        "Asegúrese de usar el formato {{NOMBRE_CAMPO}}.")

    # Guardar template en memoria
    _template_store['bytes'] = content
    _template_store['filename'] = file.filename

    print(f"[EXTRACT] '{file.filename}' -> {len(placeholders)} placeholders: {placeholders}")

    return {
        "filename": file.filename,
        "placeholders": placeholders,
        "count": len(placeholders),
    }


@app.post("/api/generate")
async def generate(request: GenerateRequest):
    """
    Genera el .docx final reemplazando placeholders con los valores provistos.
    Campos opcionales marcados como vacíos → se reemplazan por ''.
    """
    if 'bytes' not in _template_store:
        raise HTTPException(400, detail="No hay plantilla cargada. Use /api/extract primero.")

    # Construir dict de reemplazos
    replacements = {}
    for ph, val in request.values.items():
        replacements[ph] = val.strip() if val else ''

    # Opcionales vacíos → string vacío
    for ph in request.optional_empty:
        replacements[ph] = ''

    print(f"[GENERATE] Reemplazando {len(replacements)} placeholders...")
    for k, v in replacements.items():
        preview = v[:40] + ('...' if len(v) > 40 else '')
        print(f"  {{{{ {k} }}}} -> '{preview}'")

    try:
        doc = Document(io.BytesIO(_template_store['bytes']))
        doc = replace_placeholders(doc, replacements)
        
        # Validación post-generación
        remaining = check_for_remaining_placeholders(doc)
        if remaining:
            print(f"[GENERATE] ADVERTENCIA: Quedaron placeholders sin reemplazar: {remaining}")
            
    except Exception as e:
        print(f"[GENERATE] Error crítico durante la generación: {e}")
        traceback.print_exc()
        raise HTTPException(500, detail=f"Error al generar el documento: {e}")

    # Serializar a bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    original_name = _template_store.get('filename', 'contrato.docx')
    out_name = original_name.replace('.docx', '_COMPLETADO.docx')

    print(f"[GENERATE] OK -> '{out_name}'")

    return StreamingResponse(
        output,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={
            'Content-Disposition': f'attachment; filename="{out_name}"',
            'X-Unreplaced-Placeholders': ",".join(remaining) if remaining else ""
        }
    )


# ─── Static files (frontend) ─────────────────────────────────────────────────
FRONTEND_DIR = Path(__file__).parent.parent / "frontend"
if FRONTEND_DIR.exists():
    app.mount("/", StaticFiles(directory=str(FRONTEND_DIR), html=True), name="static")
